"""智能生成 API —— 文书智能编写、笔录实时整理、法条推荐、文书链批量生成、法律期限预警、文书文件上传解析、Word 导出。"""

import io
import os
import re
import tempfile
import logging

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.core.document_generator import DocumentGenerator
from app.core.file_parser import MAX_FILE_SIZE, parse_uploaded_file, validate_file
from app.core.model_manager import ModelManager
from app.core.rag_retriever import RAGRetriever
from app.core.security import sanitize_llm_input, rate_limiter
from app.models.requests import (
    DocumentGenerateRequest,
    ExtractElementsRequest,
    PolishFactRequest,
    SuggestLawsRequest,
)

logger = logging.getLogger(__name__)

# ---- 新增请求模型 ----

class DocumentChainRequest(BaseModel):
    input_text: str = Field(..., description="案情描述")
    doc_types: list[str] = Field(..., description="需要生成的多份文书类型列表")


class DeadlineCheckRequest(BaseModel):
    doc_type: str = Field(..., description="文书类型")
    fields: dict = Field(default_factory=dict, description="文书字段键值对")


class FillTemplateRequest(BaseModel):
    doc_type: str = Field(..., description="文书类型")
    fields: dict = Field(default_factory=dict, description="用户填写的字段键值对")

def _normalize_chinese_datetime(value: str) -> str:
    """将 ISO 格式时间转换为中文格式：2026-05-22 11:20 → 2026年5月22日11时20分
    也处理仅有日期的格式：2026-05-22 → 2026年5月22日
    """
    if not value:
        return value
    if "年" in value and "月" in value:
        return value
    # 带时间的完整 datetime：2026-05-22 11:20 或 2026-05-22T11:20
    m = re.match(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})[\sT](\d{1,2}):(\d{2})", value)
    if m:
        y, mo, d, h, mi = m.groups()
        return f"{y}年{int(mo)}月{int(d)}日{int(h)}时{int(mi)}分"
    # 仅有日期：2026-05-22
    m = re.match(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})$", value)
    if m:
        y, mo, d = m.groups()
        return f"{y}年{int(mo)}月{int(d)}日"
    return value


router = APIRouter(prefix="/api/generation", tags=["智能生成"])

model_manager = ModelManager()
doc_generator = DocumentGenerator(model_manager=model_manager)
retriever = RAGRetriever()


@router.post("/document")
async def generate_document(body: DocumentGenerateRequest, request: Request):
    """智能文书编写：输入案情要点 → 要素抽取 + 模板填充 + 法条推荐 + 事实润色。"""
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limiter.is_allowed(client_ip):
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")
    sanitized_input = sanitize_llm_input(body.input_text)
    result = await doc_generator.generate_document(body.doc_type, sanitized_input)
    return {
        "doc_type": body.doc_type,
        "elements": result["elements"],
        "suggested_laws": result["suggested_laws"],
        "case_nature": result["case_nature"],
        "content": result["content"],
    }


@router.post("/extract-elements")
async def extract_elements(body: ExtractElementsRequest):
    """要素抽取（不生成完整文书）。"""
    result = await doc_generator.extract_elements(body.input_text, body.doc_type)
    return {
        "elements": result.get("elements", {}),
        "suggested_laws": result.get("suggested_laws", []),
        "case_nature": result.get("case_nature", ""),
    }


@router.post("/polish-fact")
async def polish_fact(body: PolishFactRequest):
    """事实描述润色：口语化 → 规范法律文书语言。"""
    polished = await doc_generator.polish_fact(body.raw_fact, body.doc_type)
    return {"original": body.raw_fact, "polished": polished}


@router.post("/suggest-laws")
async def suggest_laws(body: SuggestLawsRequest):
    """根据案情推荐适用法条。"""
    laws = await doc_generator.suggest_laws(body.fact, body.doc_type)
    return {"suggested_laws": laws}


@router.get("/templates")
async def list_templates():
    """获取可用的文书模板列表。"""
    return {
        "templates": [
            {
                "doc_type": t.get("doc_type", ""),
                "name": t.get("name", ""),
                "description": t.get("description", ""),
            }
            for t in retriever._templates
        ],
        "total": len(retriever._templates),
    }


@router.get("/templates/{doc_type}")
async def get_template(doc_type: str):
    """获取指定类型的文书模板。"""
    template = retriever.retrieve_template(doc_type)
    if not template:
        raise HTTPException(status_code=404, detail=f"未找到类型为 {doc_type} 的文书模板")
    return template


@router.post("/document-chain")
async def generate_document_chain(body: DocumentChainRequest):
    """批量文书链生成：从同一案情描述一键生成多份关联文书。

    按办案流程顺序生成，如：受案登记表 → 传唤证 → 询问笔录 → 处罚告知 → 处罚决定。
    各文书之间共享要素抽取结果，保持事实一致性。
    """
    if len(body.doc_types) > 8:
        raise HTTPException(status_code=400, detail="单次最多生成8份文书")
    results = await doc_generator.generate_document_chain(body.input_text, body.doc_types)
    return {"results": results, "total": len(results)}


@router.post("/fill-template")
async def fill_template(body: FillTemplateRequest, request: Request):
    """手动填写模板：用户逐项填写字段 → 模板填充 + 事实润色 + 法条推荐。

    与 /document 的区别：跳过 LLM 要素抽取，直接使用用户提供的字段值填充。
    """
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limiter.is_allowed(client_ip):
        raise HTTPException(status_code=429, detail="请求过于频繁，请稍后再试")
    template = retriever.retrieve_template(body.doc_type)
    if not template:
        raise HTTPException(status_code=404, detail=f"未找到类型为 {body.doc_type} 的文书模板")

    # 清洗各字段输入
    fields = {k: sanitize_llm_input(str(v)) if v else v for k, v in body.fields.items()}
    if body.doc_type == "检查笔录":
        # 时间格式规范化：ISO → 中文
        fields["inspection_time_start"] = _normalize_chinese_datetime(fields.get("inspection_time_start", ""))
        fields["inspection_time_end"] = _normalize_chinese_datetime(fields.get("inspection_time_end", ""))
        import re
        # inspection_date: 从 inspection_time_start 提取年月日
        time_str = fields.get("inspection_time_start", "")
        date_match = re.match(r"(\d{4})年(\d{1,2})月(\d{1,2})日", time_str)
        if not date_match:
            date_match = re.match(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", time_str)
        if date_match:
            y, m, d = date_match.groups()
            fields["inspection_date"] = f"{y}年{int(m)}月{int(d)}日"
        elif time_str:
            fields["inspection_date"] = time_str

        # inspection_place: 默认为 {公安机关名称}检查室
        if not fields.get("inspection_place"):
            station = str(fields.get("police_station", "") or "")
            if station and station.lower() not in ("none", "null", ""):
                fields["inspection_place"] = f"{station}检查室"

        # inspection_target_display: 检查对象应显示为"被检查人的身体"
        target = fields.get("inspection_target", "")
        if target and not fields.get("inspection_target_display"):
            fields["inspection_target_display"] = f"{target}的身体"

        # violation_type: 去除 LLM/用户可能重复的"涉嫌"前缀（模板已含"涉嫌"）
        violation = str(fields.get("violation_type", ""))
        if violation.startswith("涉嫌"):
            violation = violation[2:]
            fields["violation_type"] = violation

        # inspection_reason: 收集{被检查人}是否有{违法性质}的证据
        fields["inspection_reason"] = f"收集{target}是否有{violation}的证据"

        # items_found: 默认"没有发现任何涉案物品"（笔录未体现则默认无）
        # 模板已含"身上"，需去除用户输入的前缀避免"身上身上"重复
        items = str(fields.get("items_found", ""))
        if items.startswith("身上"):
            items = items[2:]
        if not items or items.lower() in ("null", "none", "无", "有", "是", "否", "未发现", "未", "发现", ""):
            fields["items_found"] = "没有发现任何涉案物品"
        else:
            fields["items_found"] = items

        # injury_found: 默认"无任何损伤"（笔录未体现则默认无）
        # 模板已含"身上"，需去除用户输入的前缀避免"身上身上"重复
        injury = str(fields.get("injury_found", ""))
        if injury.startswith("身上"):
            injury = injury[2:]
        if not injury or injury.lower() in ("null", "none", "无", "有", "是", "否", "未发现", "未", "发现", ""):
            fields["injury_found"] = "无任何损伤"
        else:
            fields["injury_found"] = injury

    # 行政处罚决定书：时间格式规范化 + 自动派生字段
    if body.doc_type == "行政处罚决定书":
        import re as _re
        # ---- 时间格式规范化：ISO → 中文 ----
        fields["birth_date"] = _normalize_chinese_datetime(str(fields.get("birth_date", "") or ""))
        fields["case_occurrence_date"] = _normalize_chinese_datetime(str(fields.get("case_occurrence_date", "") or ""))
        fields["penalty_decision_date"] = _normalize_chinese_datetime(str(fields.get("penalty_decision_date", "") or ""))

        # 落款日期转换为中文大写数字格式：2026年5月15日 → 二○二六年五月十五日
        chinese_digits = {"0": "○", "1": "一", "2": "二", "3": "三", "4": "四",
                          "5": "五", "6": "六", "7": "七", "8": "八", "9": "九"}
        def _to_chinese_num(n: int) -> str:
            if n <= 9:
                return chinese_digits[str(n)]
            elif n == 10:
                return "十"
            elif n < 20:
                return "十" + chinese_digits[str(n % 10)]
            elif n % 10 == 0:
                return chinese_digits[str(n // 10)] + "十"
            else:
                return chinese_digits[str(n // 10)] + "十" + chinese_digits[str(n % 10)]

        raw_date = fields.get("penalty_decision_date", "")
        stamp_match = _re.match(r"(\d{4})年(\d{1,2})月(\d{1,2})日", raw_date)
        if stamp_match:
            y, m, d = stamp_match.groups()
            cy = "".join(chinese_digits.get(c, c) for c in y)
            cm = _to_chinese_num(int(m))
            cd = _to_chinese_num(int(d))
            fields["stamp_date"] = f"{cy}年{cm}月{cd}日"
        else:
            fields["stamp_date"] = raw_date

        # ---- 默认值 ----
        if not fields.get("work_unit") or str(fields["work_unit"]).lower() in ("null", "none", ""):
            fields["work_unit"] = "无"
        if not fields.get("criminal_record") or str(fields["criminal_record"]).lower() in ("null", "none", ""):
            fields["criminal_record"] = "无"
        # 去除末尾句号防止与模板句号重复
        cr = str(fields.get("criminal_record", ""))
        if cr.endswith("。"):
            fields["criminal_record"] = cr.rstrip("。")
        if not fields.get("native_place") or str(fields["native_place"]).lower() in ("null", "none", ""):
            fields["native_place"] = str(fields.get("household_register", ""))
        if not fields.get("household_register") or str(fields["household_register"]).lower() in ("null", "none", ""):
            fields["household_register"] = str(fields.get("party_address", ""))

        # ---- 行政复议机关 / 行政诉讼法院：从办案单位派生 ----
        dept = str(fields.get("department", ""))
        gov_match = _re.match(r"(.+?)市(.+?)公安", dept)
        county_match = _re.match(r"(.+?县).+?公安", dept)
        if not fields.get("reconsideration_org") or str(fields["reconsideration_org"]).lower() in ("null", "none", ""):
            if county_match:
                fields["reconsideration_org"] = f"{county_match.group(1)}人民政府"
            elif gov_match:
                fields["reconsideration_org"] = f"{gov_match.group(1)}市人民政府"
            elif dept:
                # fallback: 从名称中提取县级单位
                fallback = _re.match(r"(.+?[市县区])", dept)
                fields["reconsideration_org"] = f"{fallback.group(1)}人民政府" if fallback else "XX人民政府"
        if not fields.get("lawsuit_court") or str(fields["lawsuit_court"]).lower() in ("null", "none", ""):
            if county_match:
                fields["lawsuit_court"] = f"{county_match.group(1)}人民法院"
            elif gov_match:
                fields["lawsuit_court"] = f"{gov_match.group(1)}人民法院"
            elif dept:
                fallback = _re.match(r"(.+?[市县区])", dept)
                fields["lawsuit_court"] = f"{fallback.group(1)}人民法院" if fallback else "XX人民法院"

        # ---- 去除 illegal_fact 末尾的证据列举句（模板已自动添加）----
        ilf = str(fields.get("illegal_fact", ""))
        ilf = _re.sub(r"[，,。]?以上事实有[^。]+等证据证实[。]?", "", ilf)
        ilf = _re.sub(r"[，,。]?上述事实[^。]+证据[^。]证实[。]?", "", ilf)
        fields["illegal_fact"] = ilf

        # ---- 处罚依据兜底：用户未填写时从法条库自动检索 ----
        if not fields.get("penalty_basis") or str(fields["penalty_basis"]).lower() in ("null", "none", ""):
            search_text = " ".join([
                str(fields.get(k, "")) for k in ["illegal_fact", "violation_type", "case_nature"]
                if str(fields.get(k, ""))
            ])
            if search_text:
                # 中文关键词匹配（按法条 keywords 在搜索文本中的命中数打分）
                scored = []
                for law in retriever._laws:
                    score = 0
                    for kw in law.get("keywords", []):
                        if kw in search_text:
                            score += 1
                    if law.get("law_name", "") in search_text:
                        score += 3
                    if score > 0:
                        scored.append((score, law))
                scored.sort(key=lambda x: x[0], reverse=True)
                if scored:
                    top = scored[0][1]
                    fields["penalty_basis"] = f"《{top.get('law_name', '')}》第{top.get('article_number', '')}条"
                    if not fields.get("penalty_content") or str(fields["penalty_content"]).lower() in ("null", "none", ""):
                        fine = str(fields.get("fine_amount", ""))
                        if fine and fine not in ("null", "none", "", "0"):
                            fields["penalty_content"] = f"罚款{fine}元的处罚"
                        elif top.get("penalty_range"):
                            fields["penalty_content"] = top["penalty_range"]

        # ---- 执行方式和期限 ----
        if not fields.get("execution_detail") or str(fields["execution_detail"]).lower() in ("null", "none", ""):
            fields["execution_detail"] = "收到决定书之日起15日内将罚款交到指定银行财政罚没专户。"

        # ---- 文书编号（文号） ----
        if not fields.get("doc_number") or str(fields["doc_number"]).lower() in ("null", "none", ""):
            # 尝试从案件编号派生或从部门名称生成
            cn = str(fields.get("case_number", ""))
            if cn:
                fields["doc_number"] = cn
            else:
                # 格式：X公(XX)行罚决字[YYYY]NNN号
                dept_abbr = dept.replace("公安局", "").replace("分局", "")[:6] if dept else "XX"
                ds = str(fields.get("penalty_decision_date", "") or "")
                year_match = _re.search(r"(\d{4})", ds)
                year = year_match.group(1) if year_match else "YYYY"
                fields["doc_number"] = f"{dept_abbr}公行罚决字[{year}]NNN号"

    # 指认笔录：自动计算派生字段
    if body.doc_type == "指认笔录":
        # 时间格式规范化：ISO → 中文
        fields["identification_time_start"] = _normalize_chinese_datetime(fields.get("identification_time_start", ""))
        fields["identification_time_end"] = _normalize_chinese_datetime(fields.get("identification_time_end", ""))
        import re
        # identification_date: 从 identification_time_start 提取年月日
        time_str = fields.get("identification_time_start", "")
        date_match = re.match(r"(\d{4})年(\d{1,2})月(\d{1,2})日", time_str)
        if not date_match:
            date_match = re.match(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", time_str)
        if date_match:
            y, m, d = date_match.groups()
            fields["identification_date"] = f"{y}年{int(m)}月{int(d)}日"

        # case_date 格式规范化：ISO → 中文
        case_date_raw = fields.get("case_date", "")
        if case_date_raw:
            cd_match = re.match(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", case_date_raw)
            if cd_match:
                y, m, d = cd_match.groups()
                date_str = f"{y}年{int(m)}月{int(d)}日"
                if "凌晨" in case_date_raw:
                    date_str = date_str + "凌晨" + case_date_raw.split("凌晨")[-1].strip()
                fields["case_date"] = date_str

        # identification_place: 默认为案发地点（指认是在案发现场进行的）
        if not fields.get("identification_place"):
            case_loc = fields.get("case_location", "")
            if case_loc:
                fields["identification_place"] = case_loc

        # identification_purpose: 让{指认人}指认{指认对象}，锁定案件证据
        identifier = fields.get("identifier_name", "")
        target_obj = fields.get("identification_target", "")
        if identifier and target_obj and not fields.get("identification_purpose"):
            fields["identification_purpose"] = f"让{identifier}指认{target_obj}，锁定案件证据"

        # case_description: 案发经过简述（如LLM未提取，则根据其他字段自动生成）
        if not fields.get("case_description"):
            loc = fields.get("case_location", "")
            viol = fields.get("violation_type", "")
            if identifier and loc and viol:
                fields["case_description"] = f"伙同他人在{loc}{viol}了相关财物"

        # identification_confirmation: 指认确认内容（如LLM未提取，则自动生成）
        if not fields.get("identification_confirmation"):
            target = fields.get("identification_target", "")
            viol = fields.get("violation_type", "")
            if identifier and target and viol:
                fields["identification_confirmation"] = f"该{target}就是其{viol}的现场"

    # 辨认笔录：自动计算派生字段
    if body.doc_type == "辨认笔录":
        import re
        # 时间格式规范化：ISO → 中文
        fields["identification_time_start"] = _normalize_chinese_datetime(fields.get("identification_time_start", ""))
        fields["identification_time_end"] = _normalize_chinese_datetime(fields.get("identification_time_end", ""))

        # identification_place: 默认为 {公安机关名称}询问室
        if not fields.get("identification_place"):
            station = fields.get("police_station", "")
            if station:
                fields["identification_place"] = f"{station}询问室"

        # case_date 格式规范化：ISO → 中文
        case_date_raw = fields.get("case_date", "")
        if case_date_raw:
            cd_match = re.match(r"(\d{4})[/-](\d{1,2})[/-](\d{1,2})", case_date_raw)
            if cd_match:
                y, m, d = cd_match.groups()
                date_str = f"{y}年{int(m)}月{int(d)}日"
                if "凌晨" in case_date_raw:
                    date_str = date_str + "凌晨" + case_date_raw.split("凌晨")[-1].strip()
                fields["case_date"] = date_str

        identifier = fields.get("identifier_name", "")
        case_date = fields.get("case_date", "")
        case_loc = fields.get("case_location", "")
        viol = fields.get("violation_type", "")

        # identification_purpose: 让辨认人辨别确认照片中是否有在{案发日期}在{案发地点}{案件性质}的相关人员
        if not fields.get("identification_purpose"):
            if identifier and case_date and case_loc and viol:
                fields["identification_purpose"] = f"让辨认人辨别确认照片中是否有在{case_date}在{case_loc}{viol}的相关人员。"

        # identification_operation: 辨认人X将所有照片认真仔细的审视了一遍，然后指出来
        if not fields.get("identification_operation"):
            if identifier:
                fields["identification_operation"] = f"辨认人{identifier}将所有照片认真仔细的审视了一遍，然后指出来："

        # identification_conclusion: 至此，辨认结束
        if not fields.get("identification_conclusion"):
            fields["identification_conclusion"] = "至此，辨认结束。"

        # identification_confirmation: 以上笔录，X看过
        if not fields.get("identification_confirmation"):
            if identifier:
                fields["identification_confirmation"] = f"以上笔录，{identifier}看过。"

    # 现场勘查笔录：时间格式转中文
    if body.doc_type == "现场勘查笔录":
        fields["crime_scene_time_start"] = _normalize_chinese_datetime(str(fields.get("crime_scene_time_start", "") or ""))
        fields["crime_scene_time_end"] = _normalize_chinese_datetime(str(fields.get("crime_scene_time_end", "") or ""))

    # 行政处罚告知笔录：默认值 + 时间格式规范化
    if body.doc_type == "行政处罚告知笔录":
        fields["notification_time"] = _normalize_chinese_datetime(str(fields.get("notification_time", "") or ""))
        if not fields.get("legal_representative") or str(fields["legal_representative"]).lower() in ("null", "none", ""):
            fields["legal_representative"] = ""
        if not fields.get("evidence_list") or str(fields["evidence_list"]).lower() in ("null", "none", ""):
            fields["evidence_list"] = "检查笔录、违法行为人的陈述和申辩，其他"
        if not fields.get("notified_person_statement") or str(fields["notified_person_statement"]).lower() in ("null", "none", ""):
            fields["notified_person_statement"] = "我不提出陈述和申辩。"
        # 去除 proposed_penalty 末尾的"的处罚"避免与模板重复
        pp = str(fields.get("proposed_penalty", ""))
        if pp.endswith("的处罚"):
            fields["proposed_penalty"] = pp[:-3]
        # 处罚依据兜底
        if not fields.get("penalty_basis") or str(fields["penalty_basis"]).lower() in ("null", "none", ""):
            search_text = " ".join([
                str(fields.get(k, "")) for k in ["case_description", "violation_type"]
                if str(fields.get(k, ""))
            ])
            if search_text:
                scored = []
                for law in retriever._laws:
                    score = 0
                    for kw in law.get("keywords", []):
                        if kw in search_text:
                            score += 1
                    if law.get("law_name", "") in search_text:
                        score += 3
                    if score > 0:
                        scored.append((score, law))
                scored.sort(key=lambda x: x[0], reverse=True)
                if scored:
                    top = scored[0][1]
                    fields["penalty_basis"] = f"《{top.get('law_name', '')}》第{top.get('article_number', '')}条"

    template_text = template.get("template_text", "")
    content = doc_generator._fill_template(template_text, fields)

    # 推荐法条
    fact_text = " ".join(str(v) for v in fields.values() if v)
    laws = await doc_generator.suggest_laws(fact_text, body.doc_type)

    # 法律期限预警
    deadlines = doc_generator.check_legal_deadlines(fields, body.doc_type)

    return {
        "doc_type": body.doc_type,
        "elements": fields,
        "suggested_laws": [l["content"] for l in laws] if laws else [],
        "case_nature": body.fields.get("case_nature", ""),
        "content": content,
        "deadline_warnings": deadlines,
    }


@router.post("/check-deadlines")
async def check_legal_deadlines(body: DeadlineCheckRequest):
    """法律期限预警：检查当前文书涉及的法定时限是否即将到期或已超期。

    覆盖传唤时限（8/24h）、办案期限（30日）、扣押期限（30日）、告知-处罚间隔（3日）等。
    """
    warnings = doc_generator.check_legal_deadlines(body.fields, body.doc_type)
    return {"doc_type": body.doc_type, "warnings": warnings, "total": len(warnings)}


@router.post("/summarize-case-file")
async def summarize_case_file(
    file: UploadFile = File(...),
    doc_type: str = Form(""),
):
    """上传 .doc/.docx 文书 → 解析提取文本 → LLM 整理为案件摘要。

    支持格式：.doc（旧版 Word）、.docx（新版 Word）
    文件大小限制：10 MB

    返回的 summary 可直接粘贴到 AI 生成的输入框中使用。
    """
    tmp_path = ""
    try:
        # 1. 验证
        validate_file(file)

        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"文件过大，最大支持 {MAX_FILE_SIZE // 1024 // 1024} MB",
            )

        # 2. 写入临时文件
        suffix = os.path.splitext(file.filename or ".docx")[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # 3. 解析文件提取文本
        raw_text = parse_uploaded_file(tmp_path, file.filename or "unknown")
        # 移除不可打印控制字符（保留 \n \r \t），避免前端显示为问号或乱码
        raw_text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', raw_text)
        logger.info(
            "文件解析完成: filename=%s, raw_len=%d, raw_preview=%s",
            file.filename, len(raw_text), raw_text[:200],
        )

        # 4. 安全清洗（截断过长内容，移除特殊标记）
        sanitized = sanitize_llm_input(raw_text)
        logger.info(
            "安全清洗完成: sanitized_len=%d (raw=%d)",
            len(sanitized), len(raw_text),
        )

        # 5. LLM 整理为案件摘要
        result = await doc_generator.summarize_case_file_text(sanitized, doc_type)
        logger.info(
            "LLM摘要完成: filename=%s, summary_len=%d, has_warning=%s, summary_preview=%s",
            file.filename, len(result.get("summary", "")),
            bool(result.get("warning")), result.get("summary", "")[:200],
        )

        return {
            "success": True,
            "raw_text": raw_text,
            "raw_char_count": len(raw_text),
            "summary": result.get("summary", ""),
            "char_count": len(result.get("summary", "")),
            "warning": result.get("warning", ""),
        }

    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except Exception:
        logger.exception("文件处理失败: %s", file.filename)
        raise HTTPException(status_code=500, detail="文件处理失败，请联系管理员")
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except Exception:
            pass


# ── Word 导出请求模型 ──

class ExportDocxRequest(BaseModel):
    content: str = Field(..., description="文书纯文本内容")
    doc_type: str = Field("", description="文书类型（用于生成文件名）")


def _build_docx(content: str, doc_type: str = "") -> io.BytesIO:
    """将文书纯文本转换为规范公文格式的 Word 文档。"""
    is_inspection = doc_type == "检查笔录"
    is_penalty = doc_type == "行政处罚决定书"
    is_notification = doc_type == "行政处罚告知笔录"
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 默认字体 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "仿宋_GB2312"
    font.size = Pt(16)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    pf = style.paragraph_format
    pf.line_spacing = Pt(27) if is_inspection else Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    lines = content.split("\n")

    # ── 行类型前缀 ──
    FIELD_PREFIXES = (
        "检查时间", "检查地点", "检查对象",
        "检查人姓名", "见证人基本",
        "事由和目的",
        "询问时间", "询问地点", "询问人", "记录人", "被询问人",
        "案由", "案件编号",
        "时间", "地点", "侦查人员", "指认人", "指认对象", "指认目的",
        "办案人员", "辨认人姓名", "见证人姓名", "辨认对象", "辨认目的",
    )
    NARRATIVE_PREFIXES = ("过程和结果", "指认过程及结果", "辨认过程及结果")
    NOTIFICATION_FIELD_PREFIXES = ("执行告知单位", "告知人", "被告知人", "告知内容")
    NOTIFICATION_QA_PREFIXES = ("问：", "答：")
    SIGN_PREFIXES = (
        "检查人：", "见证人：", "当事人：", "被检查人：",
        "办案人：", "办案人员：", "记录人：", "审批人：", "承办人：",
        "侦查员：", "侦查人员：", "指认人：", "辨认人：",
        "被处罚人:", "被侵害人：", "接收人员：", "接收单位",
        "被告知人（及监护人）:", "被告知人（及监护人）：",
    )

    # 前两行非空行视为标题
    title_count = 0
    in_narrative = False  # 是否处于叙述段落区域（指认过程及结果 / 过程和结果 之后的内容段落也需缩进）

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            in_narrative = False
            doc.add_paragraph("")
            continue

        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing = Pt(27) if is_inspection else Pt(28)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        # ── 分类 ──
        is_sign = any(line.startswith(kw) for kw in SIGN_PREFIXES)
        is_narrative_header = any(line.startswith(kw) for kw in NARRATIVE_PREFIXES)
        is_field = not is_sign and not is_narrative_header and any(line.startswith(kw) for kw in FIELD_PREFIXES)
        is_notif_field = is_notification and not is_sign and any(line.startswith(kw) for kw in NOTIFICATION_FIELD_PREFIXES)
        is_notif_qa = is_notification and not is_sign and any(line.startswith(kw) for kw in NOTIFICATION_QA_PREFIXES)
        is_notif_date = is_notification and ("年" in line and "月" in line and "日" in line and "时" in line and "分" in line and len(line) < 20)
        is_doc_number = is_penalty and ("罚决字" in line or "行罚字" in line or "公行罚决字" in line)
        is_title = not is_sign and not is_narrative_header and not is_field and not is_doc_number and not is_notif_field and not is_notif_qa and not is_notif_date and not in_narrative and title_count < 2

        # 检查笔录的字体大小（仿宋 15pt = 小三号，其余文书 16pt = 三号）
        body_size = Pt(15) if is_inspection else Pt(16)

        if is_sign:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        elif is_doc_number:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf.first_line_indent = Pt(0)
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        elif is_notif_field:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        elif is_notif_date:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf.first_line_indent = Pt(0)
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        elif is_notif_qa:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        elif is_title:
            in_narrative = False
            title_count += 1
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Pt(0)
            run = para.add_run(line)
            if is_inspection:
                # 检查笔录：首行为单位名称（15pt，字符间加空格撑开），次行为标题（24pt 小二号）
                run.font.name = "仿宋_GB2312"
                if title_count == 1:
                    run.font.size = Pt(15)
                    # 单位名称：中文汉字之间插入一个空格撑开版面
                    spaced = "".join(ch + " " if '一' <= ch <= '鿿' else ch for ch in line)
                    run.text = spaced.rstrip()
                else:
                    run.font.size = Pt(24)
                run.bold = False
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
            else:
                if (is_penalty or is_notification) and title_count == 1:
                    # 行政处罚决定书/告知笔录：首行为公安机关名称（仿宋，居中）
                    run.font.name = "仿宋_GB2312"
                    run.font.size = Pt(16)
                    run.bold = False
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
                else:
                    run.font.name = "宋体"
                    run.font.size = Pt(22)
                    run.bold = True
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        elif is_narrative_header:
            in_narrative = True
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Pt(0) if is_inspection else Pt(32)
            if is_inspection:
                pf.line_spacing = 1.0  # 检查笔录正文单倍行距（相对值，随字号自适应）
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        elif in_narrative:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Pt(0) if is_inspection else Pt(32)
            if is_inspection:
                pf.line_spacing = 1.0  # 检查笔录正文单倍行距（相对值，随字号自适应）
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        else:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(32) if is_penalty else Pt(0)
            run = para.add_run(line)
            run.font.name = "仿宋_GB2312"
            run.font.size = body_size
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@router.post("/export-docx")
async def export_docx(body: ExportDocxRequest):
    """将生成的文书内容导出为规范公文格式的 Word 文档（.docx）。"""
    if not body.content.strip():
        raise HTTPException(status_code=400, detail="文书内容为空，无法导出")

    try:
        buf = _build_docx(body.content, body.doc_type)
    except Exception:
        logger.exception("Word 导出失败")
        raise HTTPException(status_code=500, detail="Word 文档生成失败，请联系管理员")

    safe_name = body.doc_type or "文书"
    filename = f"{safe_name}.docx"

    # 写入临时文件，用 FileResponse 正确处理中文文件名
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    try:
        tmp.write(buf.read())
        tmp_path = tmp.name
        tmp.close()
        from starlette.responses import FileResponse
        from starlette.background import BackgroundTask
        return FileResponse(
            tmp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
            background=BackgroundTask(os.unlink, tmp_path),
        )
    except Exception:
        os.unlink(tmp.name)
        raise
