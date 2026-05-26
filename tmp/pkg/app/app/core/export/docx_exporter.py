"""Word 文档导出器 —— 将文书纯文本转换为规范公文格式的 .docx 文件。

从 generation.py 中拆分出来，原先的 _build_docx 函数和导出逻辑
与 API 路由混在一起（约 200 行），现独立为单独模块。
"""

from __future__ import annotations
import io
import os
import tempfile

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from fastapi.responses import FileResponse
from starlette.background import BackgroundTask


def build_docx(content: str, doc_type: str = "") -> io.BytesIO:
    """将文书纯文本转换为规范公文格式的 Word 文档。"""
    is_inspection = doc_type == "检查笔录"
    is_penalty = doc_type == "行政处罚决定书"
    is_notification = doc_type == "行政处罚告知笔录"

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # 默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "仿宋_GB2312"
    font.size = Pt(16)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    pf = style.paragraph_format
    pf.line_spacing = Pt(27) if is_inspection else Pt(28)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    lines = content.split("\n")

    FIELD_PREFIXES = (
        "检查时间", "检查地点", "检查对象",
        "检查人姓名", "见证人基本", "事由和目的",
        "询问时间", "询问地点", "询问人", "记录人", "被询问人",
        "案由", "案件编号",
        "时间", "地点", "侦查人员", "指认人", "指认对象", "指认目的",
        "办案人员", "辨认人姓名", "见证人姓名", "辨认对象", "辨认目的",
    )
    NARRATIVE_PREFIXES = ("过程和结果", "指认过程及结果", "辨认过程及结果")
    NOTIFICATION_FIELD_PREFIXES = ("执行告知单位", "告知人", "被告知人", "告知内容")
    NOTIFICATION_QA_PREFIXES = ("问：", "答：")
    SIGN_PREFIXES = (
        "检查人：", "见证人：", "当事人：", "被检查人：",
        "办案人：", "办案人员：", "记录人：", "审批人：", "承办人：",
        "侦查员：", "侦查人员：", "指认人：", "辨认人：",
        "被处罚人:", "被侵害人：", "接收人员：", "接收单位",
        "被告知人（及监护人）:", "被告知人（及监护人）：",
    )

    title_count = 0
    in_narrative = False
    body_size = Pt(15) if is_inspection else Pt(16)

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            in_narrative = False
            doc.add_paragraph("")
            continue

        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing = Pt(27) if is_inspection else Pt(28)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        is_sign = any(line.startswith(kw) for kw in SIGN_PREFIXES)
        is_narrative_header = any(line.startswith(kw) for kw in NARRATIVE_PREFIXES)
        is_field = not is_sign and not is_narrative_header and any(
            line.startswith(kw) for kw in FIELD_PREFIXES)
        is_notif_field = is_notification and not is_sign and any(
            line.startswith(kw) for kw in NOTIFICATION_FIELD_PREFIXES)
        is_notif_qa = is_notification and not is_sign and any(
            line.startswith(kw) for kw in NOTIFICATION_QA_PREFIXES)
        is_notif_date = is_notification and (
            "年" in line and "月" in line and "日" in line
            and "时" in line and "分" in line and len(line) < 20)
        is_doc_number = is_penalty and ("罚决字" in line or "行罚字" in line)
        is_title = (
            not is_sign and not is_narrative_header and not is_field
            and not is_doc_number and not is_notif_field and not is_notif_qa
            and not is_notif_date and not in_narrative and title_count < 2)

        if is_sign:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            _add_run(para, line, body_size)
        elif is_doc_number:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf.first_line_indent = Pt(0)
            _add_run(para, line, body_size)
        elif is_notif_field or is_notif_date or is_notif_qa:
            in_narrative = False
            pf.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if is_notif_date
                else WD_ALIGN_PARAGRAPH.LEFT)
            pf.first_line_indent = Pt(0)
            _add_run(para, line, body_size)
        elif is_title:
            in_narrative = False
            title_count += 1
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Pt(0)
            if is_inspection:
                _add_run(para, _space_out_chinese(line) if title_count == 1 else line,
                        Pt(15) if title_count == 1 else Pt(24))
            elif (is_penalty or is_notification) and title_count == 1:
                _add_run(para, line, Pt(16))
            else:
                run = para.add_run(line)
                run.font.name = "宋体"
                run.font.size = Pt(22)
                run.bold = True
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        elif is_narrative_header or in_narrative:
            in_narrative = True
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf.first_line_indent = Pt(0) if is_inspection else Pt(32)
            if is_inspection:
                pf.line_spacing = 1.0
            _add_run(para, line, body_size)
        else:
            in_narrative = False
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(32) if is_penalty else Pt(0)
            _add_run(para, line, body_size)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_run(para, text: str, size: Pt) -> None:
    run = para.add_run(text)
    run.font.name = "仿宋_GB2312"
    run.font.size = size
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")


def _space_out_chinese(text: str) -> str:
    """中文字符间插入空格（检查笔录标题排版）。"""
    result = []
    for ch in text:
        result.append(ch)
        if '一' <= ch <= '鿿':
            result.append(' ')
    return ''.join(result).rstrip()


def file_response_from_docx(content: str, doc_type: str) -> FileResponse:
    """从文书内容生成临时文件并返回 FileResponse。"""
    buf = build_docx(content, doc_type)
    safe_name = doc_type or "文书"
    filename = f"{safe_name}.docx"

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    try:
        tmp.write(buf.read())
        tmp_path = tmp.name
        tmp.close()
        return FileResponse(
            tmp_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
            background=BackgroundTask(os.unlink, tmp_path),
        )
    except Exception:
        os.unlink(tmp.name)
        raise
